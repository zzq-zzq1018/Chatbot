from flask import Flask, render_template, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, date
import requests
import json
import time
import threading
from collections import defaultdict
import io
import re

app = Flask(__name__)

# 数据库配置
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///chat.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)


# 对话日志表（增加风险字段）
class ChatLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), default="user")
    user_msg = db.Column(db.Text, nullable=False)
    bot_msg = db.Column(db.Text, nullable=False)
    context = db.Column(db.Text, default="[]")
    scene = db.Column(db.String(20), default="auto", index=True)
    is_pro = db.Column(db.Boolean, default=False)  # 专业模式标记
    risk_level = db.Column(db.String(20), default="low")  # 风险等级: low/medium/high
    risk_alert = db.Column(db.Text, default="")  # 风险提醒内容
    create_time = db.Column(db.DateTime, default=datetime.now, index=True)


# 压力测试结果存储
test_results = {
    "total_requests": 0,
    "success_requests": 0,
    "fail_requests": 0,
    "avg_response_time": 0.0,
    "response_times": [],
    "error_details": defaultdict(int)
}

# ====================== 风险评估关键词配置 ======================
# 民事高风险关键词
CIVIL_HIGH_RISK_KEYWORDS = [
    "家暴", "家庭暴力", "殴打", "人身伤害", "重伤", "打死", "打伤",
    "大额资金", "百万", "千万", "巨额", "全部积蓄", "养老钱", "救命钱",
    "紧急", "立刻", "马上", "今天", "今晚", "正在", "威胁", "恐吓"
]

# 刑事高风险关键词
CRIMINAL_HIGH_RISK_KEYWORDS = [
    "拘留", "刑事拘留", "逮捕", "被抓", "带走", "关押", "看守所",
    "在逃", "通缉", "跑路", "躲避", "生命安全", "生命危险", "杀人", "故意杀人",
    "判刑", "坐牢", "有期徒刑", "无期徒刑"
]

# 劳动高风险关键词
LABOR_HIGH_RISK_KEYWORDS = [
    "工伤", "伤残", "死亡", "工亡", "危险作业", "安全", "爆炸", "坠落",
    "恶意欠薪", "跑路", "失联", "老板跑了", "公司倒闭", "被迫离职", "逼迫", "威胁辞职"
]

# 民事中风险关键词
CIVIL_MEDIUM_RISK_KEYWORDS = [
    "起诉", "诉讼", "法院", "开庭", "判决", "执行", "查封", "冻结",
    "离婚", "抚养权", "财产分割", "债务", "欠款", "逾期"
]

# 劳动中风险关键词
LABOR_MEDIUM_RISK_KEYWORDS = [
    "拖欠工资", "不发工资", "克扣", "辞退", "解雇", "开除", "裁员",
    "没签合同", "未签合同", "社保", "五险一金"
]

# 刑事中风险关键词
CRIMINAL_MEDIUM_RISK_KEYWORDS = [
    "诈骗", "盗窃", "抢劫", "故意伤害", "轻伤", "立案", "报案",
    "取保", "取保候审", "监视居住"
]

# ====================== 普通模式提示词 ======================
LEGAL_PROMPT_TEMPLATES = {
    "civil": """
你是专业的民事纠纷法律咨询助手，严格遵守以下规则：
1. 仅回答民间借贷、婚姻家庭、房产纠纷、欠钱不还、借款等民事相关问题，非民事法律问题直接拒绝；
2. 法律分析仅作参考，必须提示"不可替代律师专业意见，复杂案件建议咨询执业律师"；
3. 回答控制在150字左右，简洁清晰、专业易懂；
4. 多轮对话需记住上下文，衔接用户历史问题；
5. 涉及诉讼时效、证据收集等关键信息，需明确提示；
6. 回答完后，根据用户当前问题，自动推荐3个最相关的可继续提问的问题。
7. 格式要求：回答内容 + 换行 + 💡你可以继续问：问题1？问题2？问题3？
""",
    "criminal": """
你是专业的刑事辩护法律咨询助手，严格遵守以下规则：
1. 仅回答盗窃、故意伤害、诈骗等刑事相关问题，非刑事法律问题直接拒绝；
2. 法律分析仅作参考，必须提示"不可替代律师辩护意见，被采取强制措施请立即委托律师"；
3. 回答控制在150字左右，专业严谨、条理清晰；
4. 多轮对话需记住用户提及的案件情节、当事人身份等上下文信息；
5. 涉及刑事拘留、逮捕等紧急情况，强制建议立即联系刑事辩护律师；
6. 回答完后，根据用户当前问题，自动推荐3个最相关的可继续提问的问题。
7. 格式要求：回答内容 + 换行 + 💡你可以继续问：问题1？问题2？问题3？
""",
    "labor": """
你是专业的劳动仲裁法律咨询助手，严格遵守以下规则：
1. 仅回答劳动合同、工资薪酬、工伤赔偿等劳动相关问题，非劳动法律问题直接拒绝；
2. 法律分析仅作参考，必须提示"不可替代劳动仲裁委/法院裁决，争议建议申请劳动仲裁"；
3. 回答控制在150字左右，通俗实用、步骤清晰；
4. 多轮对话需记住用户的工作年限、薪资、纠纷类型等上下文信息；
5. 涉及拖欠工资、违法辞退等紧急情况，明确提示仲裁时效和维权渠道；
6. 回答完后，根据用户当前问题，自动推荐3个最相关的可继续提问的问题。
7. 格式要求：回答内容 + 换行 + 💡你可以继续问：问题1？问题2？问题3？
"""
}

# ====================== 专业模式提示词（面向律师/法律从业者） ======================
PRO_PROMPT_TEMPLATES = {
    "civil": """
你是资深民事诉讼律师，正在为同行（律师/法务/法律工作者）提供专业支持。

【回答风格要求】
- 直接使用法言法语，无需通俗解释，无需免责提示
- 结论先行，层层递进
- 使用"构成要件""举证责任""裁判规则""诉讼策略"等专业表述
- 以"律师对律师"的口吻交流

【回答格式要求】请严格按照以下结构组织回答：

**一、法律依据与构成要件**
- 引用具体法律条文（民法典第X条、司法解释第X条）
- 列明请求权基础或抗辩事由的构成要件

**二、举证责任分配与证据清单**
- 各方当事人的举证责任
- 需要收集的关键证据类型及证明目的
- 证据瑕疵风险提示

**三、诉讼策略与风险要点**
- 起诉/答辩的思路选择
- 常见裁判观点及类案检索方向
- 程序性风险（时效、管辖、保全等）

**四、类案裁判观点**
- 最高院/地方高院的典型案例裁判要旨
- 司法实践中的主流观点与分歧

**五、实操建议**
- 文书撰写要点
- 庭审发问/质证要点
- 调解/和解策略

**六、专业延伸（3个问题）**
格式：💡专业延伸：问题1？问题2？问题3？

【核心原则】
- 不解释基础概念（律师默认掌握）
- 不提供当事人视角的"维权建议"
- 专注于法律技术分析
""",
    "criminal": """
你是资深刑事辩护律师，正在为同行（律师/法务/法律工作者）提供专业支持。

【回答风格要求】
- 直接使用法言法语，无需通俗解释
- 使用"犯罪构成""证据链""辩护方向""量刑情节"等专业表述
- 以"律师对律师"的口吻交流

【回答格式要求】请严格按照以下结构组织回答：

**一、罪名构成要件与法律依据**
- 引用刑法条文及相关司法解释
- 四要件/三阶层分析框架

**二、证据审查与质证要点**
- 定罪证据体系分析
- 关键证据的质证思路（非法证据排除、证据链断裂点）
- 存疑有利于被告原则的适用空间

**三、辩护策略选择**
- 无罪辩护/罪轻辩护/程序辩护的适用场景
- 量刑情节辩护（自首、立功、认罪认罚、退赔谅解）
- 变更强制措施（取保候审）的时机与策略

**四、类案检索与裁判观点**
- 类似案件的量刑区间及裁判规则
- 本地法院的裁量倾向

**五、风险提示与应对**
- 诉讼阶段风险（侦查/审查起诉/审判）
- 当事人权益保护的实操要点

**六、专业延伸（3个问题）**
格式：💡专业延伸：问题1？问题2？问题3？

【核心原则】
- 不解释基础刑法概念（律师默认掌握）
- 专注于辩护技术和诉讼策略
""",
    "labor": """
你是资深劳动法律师（劳动仲裁专家），正在为同行（律师/法务/HR/法律工作者）提供专业支持。

【回答风格要求】
- 直接使用法言法语，无需通俗解释
- 使用"违法解除""赔偿基数""举证倒置""裁审口径"等专业表述
- 以"律师对律师"的口吻交流

【回答格式要求】请严格按照以下结构组织回答：

**一、法律依据与请求权基础**
- 引用《劳动合同法》《劳动争议调解仲裁法》具体条文
- 列明适用的法律要件

**二、证据清单与举证策略**
- 劳动者需要提交的证据及证明目的
- 用人单位可采用的抗辩证据
- 举证责任倒置规则的运用

**三、仲裁/诉讼策略分析**
- 仲裁请求的设计（列明具体请求事项）
- 赔偿计算标准（N/N+1/2N的适用条件及计算方式）
- 抗辩思路及常见争议焦点

**四、裁审口径与类案参考**
- 本地仲裁委/法院的裁判倾向
- 最高院指导案例或典型案例的裁判要旨

**五、风险要点与实操建议**
- 时效风险（仲裁时效1年，起算点认定）
- 程序风险（终局裁决的撤销条件）
- 执行风险及应对

**六、专业延伸（3个问题）**
格式：💡专业延伸：问题1？问题2？问题3？

【核心原则】
- 不解释基础劳动法概念
- 专注于仲裁/诉讼技术、赔偿计算、程序策略
"""
}

# ====================== 自动场景识别关键词 ======================
CIVIL_KEYWORDS = ["借钱", "欠钱", "借条", "转账", "债务", "还款", "彩礼", "离婚", "抚养权", "房产", "房东", "租客",
                  "押金", "交通事故", "合同", "违约", "买卖", "继承", "遗嘱", "赠与", "不当得利"]
CRIMINAL_KEYWORDS = ["盗窃", "诈骗", "轻伤", "拘留", "逮捕", "取保", "缓刑", "坐牢", "立案", "罪名", "故意伤害", "抢劫",
                     "强奸", "贪污", "受贿", "帮信", "掩饰隐瞒", "取保候审", "监视居住"]
LABOR_KEYWORDS = ["工资", "辞退", "拖欠", "加班", "社保", "工伤", "劳动合同", "仲裁", "工地", "赔偿金", "经济补偿",
                  "N+1", "2N", "违法解除", "竞业限制"]


# ====================== DeepSeek 配置 ======================
DEEPSEEK_API_KEY = "sk-efb0b7c994c84df6aab93c6a66a1ad1f"
DEEPSEEK_URL = "https://api.deepseek.com/v1/chat/completions"


# ====================== 风险评估核心函数 ======================
def detect_risk_level(user_msg, scene):
    """
    风险等级判定
    返回: (risk_level, risk_alert, risk_reason)
    risk_level: 'high', 'medium', 'low'
    risk_alert: 风险提示文本
    risk_reason: 触发原因
    """
    msg_lower = user_msg.lower()

    # 高优先级：先检查高风险关键词
    if scene == "civil":
        for kw in CIVIL_HIGH_RISK_KEYWORDS:
            if kw in msg_lower:
                risk_alert = "⚠️ 【高风险提醒】您描述的情况涉及人身安全或重大财产风险。建议：1) 首先确保自身安全，必要时拨打110；2) 尽快咨询执业律师；3) 保存所有相关证据。"
                return "high", risk_alert, f"触发民事高风险关键词: {kw}"
    elif scene == "criminal":
        for kw in CRIMINAL_HIGH_RISK_KEYWORDS:
            if kw in msg_lower:
                risk_alert = "⚠️ 【刑事高风险提醒】您描述的情况涉及刑事强制措施或人身安全。强烈建议：1) 立即委托专业刑事律师介入；2) 配合调查但有权保持沉默；3) 不要自行与对方私下解决。"
                return "high", risk_alert, f"触发刑事高风险关键词: {kw}"
    elif scene == "labor":
        for kw in LABOR_HIGH_RISK_KEYWORDS:
            if kw in msg_lower:
                risk_alert = "⚠️ 【劳动高风险提醒】您描述的情况涉及工伤、危险作业或恶意欠薪。建议：1) 第一时间就医并保留病历；2) 收集证据（工资单、合同、聊天记录）；3) 尽快申请劳动仲裁或投诉劳动监察大队。"
                return "high", risk_alert, f"触发劳动高风险关键词: {kw}"

    # 检查跨场景高风险（人身安全类）
    safety_keywords = ["生命危险", "打死", "杀", "逃跑", "紧急"]
    for kw in safety_keywords:
        if kw in msg_lower:
            risk_alert = "⚠️ 【紧急高风险提醒】您描述的情况可能涉及人身安全威胁！请立即拨打110报警，并确保自身安全。本平台建议尽快联系当地律师获得紧急法律协助。"
            return "high", risk_alert, f"触发安全高风险关键词: {kw}"

    # 检查中风险
    if scene == "civil":
        for kw in CIVIL_MEDIUM_RISK_KEYWORDS:
            if kw in msg_lower:
                risk_alert = "🔔 【中风险提醒】您的情况需要尽快处理。建议：1) 咨询专业律师；2) 收集相关证据；3) 注意诉讼时效。"
                return "medium", risk_alert, f"触发民事中风险关键词: {kw}"
    elif scene == "criminal":
        for kw in CRIMINAL_MEDIUM_RISK_KEYWORDS:
            if kw in msg_lower:
                risk_alert = "🔔 【中风险提醒】您的情况涉及刑事法律风险。建议：1) 咨询专业刑事律师；2) 不要自行处理敏感事项；3) 保存所有相关证据。"
                return "medium", risk_alert, f"触发刑事中风险关键词: {kw}"
    elif scene == "labor":
        for kw in LABOR_MEDIUM_RISK_KEYWORDS:
            if kw in msg_lower:
                risk_alert = "🔔 【中风险提醒】您的情况建议尽快维权。建议：1) 收集工资单、合同、考勤记录；2) 向劳动监察大队投诉；3) 考虑申请劳动仲裁。"
                return "medium", risk_alert, f"触发劳动中风险关键词: {kw}"

    # 低风险（默认）
    risk_alert = "ℹ️ 【低风险提醒】您的情况属于一般法律咨询。以上分析仅供参考，如情况复杂建议咨询专业律师。"
    return "low", risk_alert, "未触发高风险/中风险关键词"


def merge_risk_with_reply(bot_reply, risk_level, risk_alert):
    """
    将风险提示融合到机器人回答中
    """
    if risk_level == "high":
        return f"{risk_alert}\n\n---\n\n{bot_reply}"
    elif risk_level == "medium":
        return f"{bot_reply}\n\n---\n\n{risk_alert}"
    else:
        return f"{bot_reply}\n\n{risk_alert}"


# ====================== DeepSeek 调用 ======================
def auto_detect_scene(user_msg):
    msg = user_msg.lower()
    labor_score = sum(1 for kw in LABOR_KEYWORDS if kw in msg)
    criminal_score = sum(1 for kw in CRIMINAL_KEYWORDS if kw in msg)
    civil_score = sum(1 for kw in CIVIL_KEYWORDS if kw in msg)

    if labor_score >= 2 and labor_score > max(civil_score, criminal_score):
        return "labor"
    elif criminal_score >= 2 and criminal_score > max(civil_score, labor_score):
        return "criminal"
    elif civil_score >= 2:
        return "civil"
    elif labor_score >= criminal_score and labor_score >= civil_score:
        return "labor" if labor_score > 0 else "civil"
    elif criminal_score >= civil_score:
        return "criminal"
    else:
        return "civil"


def call_deepseek(messages, scene="civil", is_pro=False, retry_times=2):
    if is_pro:
        system_prompt = PRO_PROMPT_TEMPLATES.get(scene, PRO_PROMPT_TEMPLATES["civil"])
    else:
        system_prompt = LEGAL_PROMPT_TEMPLATES.get(scene, LEGAL_PROMPT_TEMPLATES["civil"])

    messages[0]["content"] = system_prompt
    max_context = 12 if is_pro else 8
    messages = messages[:max_context * 2 + 1]

    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "deepseek-chat",
        "messages": messages,
        "temperature": 0.15 if is_pro else 0.25,
        "stream": False
    }

    for retry in range(retry_times + 1):
        try:
            start_time = time.time()
            resp = requests.post(DEEPSEEK_URL, headers=headers, json=payload, timeout=35)
            end_time = time.time()
            response_time = end_time - start_time

            if resp.status_code == 200:
                data = resp.json()
                return {
                    "reply": data["choices"][0]["message"]["content"],
                    "success": True,
                    "response_time": response_time
                }

            if resp.status_code in (429, 500, 502, 503, 504):
                if retry < retry_times:
                    time.sleep(0.8 * (retry + 1))
                    continue
                else:
                    raise Exception(f"API服务异常 {resp.status_code}")
            else:
                raise Exception(f"API返回错误 {resp.status_code}")

        except requests.exceptions.ConnectTimeout:
            if retry < retry_times:
                time.sleep(0.8 * (retry + 1))
                continue
            error_msg = "连接超时"
        except requests.exceptions.ReadTimeout:
            if retry < retry_times:
                time.sleep(0.8 * (retry + 1))
                continue
            error_msg = "响应超时"
        except Exception as e:
            error_msg = str(e)
            break

    return {
        "reply": "服务暂时不可用，请稍后重试。",
        "success": False,
        "response_time": 0,
        "error": error_msg
    }


# ====================== AI 生成法律文书 ======================
@app.route('/generate-document', methods=['POST'])
def generate_document():
    """根据对话历史生成法律文书草稿"""
    data = request.get_json()
    user_msg = data.get("user_msg", "")
    context = data.get("context", [])
    doc_type = data.get("doc_type", "complaint")

    # 收集所有对话内容
    all_dialogues = []
    for item in context[-15:]:
        if item.get("user"):
            all_dialogues.append(f"用户：{item['user']}")
        if item.get("bot"):
            all_dialogues.append(f"助手：{item['bot']}")
    all_dialogues.append(f"用户：{user_msg}")
    all_text = "\n".join(all_dialogues)

    doc_prompts = {
        "complaint": f"""你是一位资深法律文书专家。请根据以下用户描述的法律纠纷，生成一份完整的《民事起诉状》。

## 用户描述的情况：
{all_text}

## 要求：
1. 严格按照起诉状格式，包含：原告信息、被告信息、诉讼请求、事实与理由、此致、具状人
2. 原告/被告信息中缺失的内容用「待补充」标注
3. 诉讼请求必须具体、明确、可执行
4. 事实与理由需逻辑清晰，引用适当法律原则
5. 直接输出起诉状正文，不要加额外解释

请输出：""",
        "defense": f"""你是一位资深法律文书专家。请根据以下用户描述的案件情况，生成一份完整的《民事答辩状》。

## 用户描述的情况：
{all_text}

## 要求：
1. 严格按照答辩状格式，包含：答辩人信息、答辩请求、事实与理由
2. 针对原告可能的指控进行逐条辩驳
3. 答辩请求明确具体
4. 直接输出答辩状正文，不要加额外解释

请输出：""",
        "labor_application": f"""你是一位资深劳动法律师。请根据以下用户描述的劳动纠纷，生成一份完整的《劳动仲裁申请书》。

## 用户描述的情况：
{all_text}

## 要求：
1. 包含：申请人信息、被申请人信息、仲裁请求、事实与理由
2. 仲裁请求包括：工资、经济补偿金、双倍工资差额、加班费、年假折算等（根据案情）
3. 引用《劳动合同法》相关条款
4. 直接输出申请书正文，不要加额外解释

请输出：""",
        "contract_clause": f"""你是一位合同审查专家。请根据以下用户的需求，生成针对性的《合同条款建议》。

## 用户想修改/草拟的合同内容：
{all_text}

## 要求：
1. 输出3-5条核心合同条款建议
2. 每条条款包含：条款名称、条款内容
3. 语言专业、可直接用于合同
4. 直接输出建议正文

请输出："""
    }

    prompt = doc_prompts.get(doc_type, doc_prompts["complaint"])

    messages = [
        {"role": "system", "content": "你是专业的法律文书生成助手，擅长根据案情生成结构完整、格式规范的诉讼文书。"},
        {"role": "user", "content": prompt}
    ]

    result = call_deepseek(messages, scene="civil", is_pro=True)

    if result["success"]:
        # 去除「专业延伸」及后面的内容
        clean_doc = result["reply"]

        # 【新增】去除 Markdown 标题符号（# 号）
        clean_doc = re.sub(r'^#{1,6}\s+', '', clean_doc, flags=re.MULTILINE)

        # 【新增】去除列表符号（行首的 - 或 • 或数字加点的格式，保留内容）
        clean_doc = re.sub(r'^[\s]*[-•]\s+', '', clean_doc, flags=re.MULTILINE)
        # 匹配并删除 💡专业延伸 及其后的所有内容
        clean_doc = re.sub(r'\n*\s*💡专业延伸[：:].*$', '', clean_doc, flags=re.DOTALL)

        clean_doc = clean_doc.replace('*', '')  # 删除所有星号
        doc_names = {
            "complaint": "民事起诉状",
            "defense": "民事答辩状",
            "labor_application": "劳动仲裁申请书",
            "contract_clause": "合同条款建议"
        }
        filename = f"{doc_names.get(doc_type, '法律文书')}_{datetime.now().strftime('%Y%m%d')}.txt"
        return jsonify({
            "success": True,
            "document": clean_doc,
            "filename": filename,
            "response_time": result.get("response_time", 0)
        })
    else:
        return jsonify({
            "success": False,
            "error": result.get("error", "生成失败，请稍后重试")
        }), 500


@app.route('/export-document', methods=['POST'])
def export_document():
    """将生成的文书导出为可下载文件"""
    data = request.get_json()
    content = data.get("content", "")
    filename = data.get("filename", "法律文书.txt")
    file_format = data.get("format", "txt")

    if not content:
        return jsonify({"error": "内容为空"}), 400

    try:
        if file_format == "docx":
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # 设置页面边距
            section = doc.sections[0]
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

            # 添加标题
            title = doc.add_heading(filename.replace('.docx', ''), level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 按段落添加内容
            for line in content.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    # 设置字体大小
                    if p.runs:
                        p.runs[0].font.size = Pt(11)

            # 保存到内存
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            return send_file(
                file_stream,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            # txt 格式
            file_stream = io.BytesIO(content.encode('utf-8'))
            return send_file(
                file_stream,
                as_attachment=True,
                download_name=filename.replace('.docx', '.txt'),
                mimetype='text/plain; charset=utf-8'
            )
    except Exception as e:
        return jsonify({"error": f"导出失败: {str(e)}"}), 500

# ====================== 路由 ======================
@app.route('/')
def index():
    return render_template('index.html')


# 普通模式聊天（集成风险评估）
@app.route('/chat', methods=['POST'])
def chat():
    data = request.get_json()
    user_msg = data.get("msg", "").strip()
    context = data.get("context", [])

    legal_keywords = [
        "法律", "纠纷", "合同", "仲裁", "诉讼", "罪名", "赔偿", "工资", "工伤", "婚姻", "房产", "借贷", "盗窃", "追尾",
        "诈骗", "借钱", "欠钱", "借条", "转账", "家暴", "起诉", "立案", "法院", "债务", "还款", "借款", "维权",
        "证据", "时效", "失信", "离婚", "出轨", "抚养权", "彩礼", "辞退", "拖欠", "加班", "社保", "轻伤", "赠与",
        "拘留", "逮捕", "取保", "缓刑", "房东", "租客", "押金", "交通事故", "跑路", "复婚", "再婚", "分居", "同居",
        "婚内财产", "婚前财产", "共同财产", "个人财产", "分家析产", "继承", "遗产", "遗嘱", "赡养", "抚养费", "探视权",
        "丧偶", "重婚", "冷暴力", "净身出户", "高利贷", "砍头息", "逾期", "滞纳金", "违约金", "担保", "抵押", "质押",
        "连带保证", "担保人", "追偿", "流水", "不当得利", "非法集资", "劳动合同", "试用期", "赔偿金", "经济补偿金",
        "N+1", "未签合同", "双倍工资", "旷工", "调岗", "降薪", "职业病", "年假", "劳务派遣", "竞业限制", "保密协议",
        "离职证明", "产权", "过户", "房产证", "不动产", "定金", "订金", "一房二卖", "物业", "漏水", "违建", "续租",
        "转租", "自首", "立功", "累犯", "共犯", "治安处罚", "管制", "拘役", "无期徒刑", "帮信", "掩饰隐瞒", "追赃",
        "人身损害", "精神损害赔偿", "伤残鉴定", "护理费", "营养费", "误工费", "丧葬费", "死亡赔偿金", "侮辱", "诽谤",
        "名誉权", "隐私权", "肖像权", "调解书", "判决书", "强制执行", "查封", "冻结", "保全", "反诉", "上诉", "再审",
        "传票", "律师函", "公证", "违约", "不可抗力", "买卖合同", "合伙", "交强险", "三者险", "全责", "主责", "同责",
        "次责", "无责", "伤残", "理赔", "定损", "酒驾", "无证驾驶", "打官司", "讨薪", "私了", "和解", "协商", "追责",
        "管辖权异议", "举证期限", "财产保全", "先予执行", "再审申请","隐匿转移财产"
        # 刑事诉讼与治安
        "取保候审", "监视居住", "逮捕条件", "拘留期限", "刑事和解", "认罪认罚", "从轻处罚", "减轻处罚", "缓刑条件",
        "假释", "减刑", "刑满释放", "前科", "犯罪记录封存", "公诉", "自诉", "受害人谅解", "被害人陈述", "讯问",
        "询问", "搜查", "扣押", "监外执行", "社区矫正",
        # 民事侵权与人身损害
        "过失相抵", "因果关系", "共同侵权", "雇主责任", "产品责任", "医疗损害", "环境污染", "高度危险作业", "动物致害",
        "建筑物倒塌", "高空抛物", "见义勇为", "自甘风险", "自助行为", "精神抚慰金", "后续治疗费", "康复费", "残疾辅助器具费",
        # 婚姻家庭与继承
        "事实婚姻", "同居析产", "非婚生子女", "亲子鉴定", "收养", "继父母子女", "扶养", "扶养费", "家庭暴力告诫书",
        "人身安全保护令", "夫妻共同债务", "个人债务", "离婚冷静期", "诉讼离婚", "协议离婚", "抚育费", "抚育费变更",
        "遗产分割", "遗赠扶养协议", "代位继承", "转继承", "必留份", "公证遗嘱", "自书遗嘱", "打印遗嘱", "录音录像遗嘱", "口头遗嘱",
        # 劳动与人事
        "工伤认定", "劳动能力鉴定", "停工留薪期", "医疗期", "病假工资", "最低工资", "加班费计算", "不定时工作制","砸伤",
        "综合计算工时", "违法解除", "继续履行劳动合同", "代通知金", "未休年假工资", "高温津贴", "女职工保护", "孕期解雇",
        "产假", "哺乳假", "工伤复发", "视同工伤",
        # 合同与债权债务
        "缔约过失", "格式条款", "显失公平", "重大误解", "欺诈撤销", "情势变更", "合同解除", "合同终止", "继续履行",
        "定金罚则", "预付款", "服务期", "质量保证金", "保函", "应收账款质押", "债务加入", "债务转移", "债权转让",
        "代位权", "撤销权", "第三人履行", "不安抗辩", "同时履行抗辩",
        # 房产物业与相邻关系
        "预告登记", "抵押权登记", "居住权", "公房承租", "经济适用房", "安置房", "小产权房", "违法强拆", "征收补偿",
        "拆迁安置", "物业费", "业主大会", "业委会", "专项维修资金", "采光权", "通风权", "噪音扰民", "邻里排水",
        "共有部分", "专有部分",
        # 交通事故与保险
        "代位求偿", "车上人员险", "不计免赔", "免责条款", "逃逸", "顶包", "实习期上高速", "非机动车与行人", "车损险",
        "划痕险", "涉水险", "停运损失", "替代性交通工具费", "评残", "重新鉴定", "抢救费用垫付", "道路救助基金","斑马线"
        # 执行与程序
        "执行异议", "执行和解", "执行担保", "执行移送", "终结本次执行", "纳入失信名单", "限制高消费", "限制出境",
        "司法拘留", "罚款", "司法拍卖", "变卖", "以物抵债", "参与分配", "执行回转", "支付令", "公示催告", "管辖权",
        "指定管辖", "移送管辖", "回避", "诉讼时效中断", "诉讼时效中止",
        # 行政与公法
        "行政复议", "行政诉讼", "国家赔偿", "行政强制", "行政处罚", "行政拘留", "吊销驾照", "责令停产停业", "征收决定",
        "政府信息公开", "信访", "行政复议前置", "复议终局",
        # 常见罪名（补充）
        "危险驾驶", "交通肇事", "故意伤害", "寻衅滋事", "开设赌场", "赌博罪", "组织卖淫", "强迫卖淫", "贩卖毒品",
        "容留他人吸毒", "非法拘禁", "绑架", "抢劫", "抢夺", "敲诈勒索", "职务侵占", "挪用资金", "非国家工作人员受贿",
        "行贿", "受贿", "滥用职权", "玩忽职守",
        # 网络与数据法律
        "网络侵权", "隐私泄露", "数据安全", "个人信息保护", "人脸识别", "网络谣言", "网络暴力", "虚拟财产", "游戏账号",
        "直播打赏", "电商平台责任", "七天无理由退货", "大数据杀熟",
        # 常用法律术语（短词）
        "法条", "司法解释", "指导案例", "判例", "举证责任", "质证", "认证", "庭前会议", "庭审", "宣判", "送达",
        "上诉状", "答辩状", "起诉状", "代理词", "辩护词", "法律援助", "司法救助"
    ]

    if not any(keyword in user_msg for keyword in legal_keywords):
        return jsonify({
            "reply": "抱歉，我仅能回答法律相关问题，请提出具体的法律咨询问题。",
            "context": context,
            "success": False,
            "response_time": 0
        })

    if not user_msg:
        return jsonify({"reply": "请输入法律问题", "context": [], "success": False})

    scene = auto_detect_scene(user_msg)

    risk_level, risk_alert, risk_reason = detect_risk_level(user_msg, scene)

    messages = [{"role": "system", "content": ""}]
    for item in context[-10:]:
        messages.append({"role": "user", "content": item["user"]})
        messages.append({"role": "assistant", "content": item["bot"]})
    messages.append({"role": "user", "content": user_msg})

    result = call_deepseek(messages, scene, is_pro=False)
    bot_reply_raw = result["reply"]
    bot_reply = merge_risk_with_reply(bot_reply_raw, risk_level, risk_alert)

    new_context = context[-10:]
    new_context.append({"user": user_msg, "bot": bot_reply})

    try:
        log = ChatLog(
            user_msg=user_msg,
            bot_msg=bot_reply,
            context=json.dumps(new_context, ensure_ascii=False),
            scene=scene,
            is_pro=False,
            risk_level=risk_level,
            risk_alert=risk_alert[:500]
        )
        db.session.add(log)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"数据库保存失败: {e}")

    return jsonify({
        "reply": bot_reply,
        "context": new_context,
        "success": result["success"],
        "response_time": result.get("response_time", 0),
        "risk_level": risk_level,
        "risk_alert": risk_alert,
        "scene": scene
    })


# 专业模式聊天（集成风险评估）
@app.route('/chat-pro', methods=['POST'])
def chat_pro():
    data = request.get_json()
    user_msg = data.get("msg", "").strip()
    context = data.get("context", [])

    if not user_msg:
        return jsonify({"reply": "请输入专业法律问题", "context": [], "success": False})

    scene = auto_detect_scene(user_msg)

    risk_level, risk_alert, risk_reason = detect_risk_level(user_msg, scene)
    if risk_level == "high":
        risk_alert = "⚠️ 高风险案件，建议立即启动应急法律程序。"
    elif risk_level == "medium":
        risk_alert = "🔔 中风险案件，建议尽快安排法律行动。"
    else:
        risk_alert = "ℹ️ 一般咨询，标准处理流程。"

    messages = [{"role": "system", "content": ""}]
    for item in context[-12:]:
        messages.append({"role": "user", "content": item["user"]})
        messages.append({"role": "assistant", "content": item["bot"]})
    messages.append({"role": "user", "content": user_msg})

    result = call_deepseek(messages, scene, is_pro=True)
    bot_reply_raw = result["reply"]
    bot_reply = merge_risk_with_reply(bot_reply_raw, risk_level, risk_alert)

    new_context = context[-12:]
    new_context.append({"user": user_msg, "bot": bot_reply})

    try:
        log = ChatLog(
            user_msg=user_msg,
            bot_msg=bot_reply,
            context=json.dumps(new_context, ensure_ascii=False),
            scene=scene,
            is_pro=True,
            risk_level=risk_level,
            risk_alert=risk_alert[:500]
        )
        db.session.add(log)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"数据库保存失败: {e}")

    return jsonify({
        "reply": bot_reply,
        "context": new_context,
        "success": result["success"],
        "response_time": result.get("response_time", 0),
        "risk_level": risk_level,
        "risk_alert": risk_alert,
        "scene": scene
    })


# 专业模式统计
@app.route('/pro-stat')
def pro_stat():
    today = date.today()
    total = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today).count()
    pro = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.is_pro == True).count()

    civil = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.scene == "civil").count()
    labor = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.scene == "labor").count()
    criminal = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.scene == "criminal").count()

    high_risk = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.risk_level == "high").count()
    medium_risk = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today,
                                       ChatLog.risk_level == "medium").count()

    all_cnt = civil + labor + criminal
    civil_pct = round(civil / all_cnt * 100, 1) if all_cnt else 0
    labor_pct = round(labor / all_cnt * 100, 1) if all_cnt else 0
    criminal_pct = round(criminal / all_cnt * 100, 1) if all_cnt else 0

    return jsonify({
        "today_total": total,
        "today_pro": pro,
        "civil_pct": civil_pct,
        "labor_pct": labor_pct,
        "criminal_pct": criminal_pct,
        "high_risk": high_risk,
        "medium_risk": medium_risk
    })


@app.route('/risk-logs')
def risk_logs():
    logs = ChatLog.query.filter(ChatLog.risk_level.in_(['high', 'medium'])).order_by(ChatLog.create_time.desc()).limit(
        50).all()
    result = []
    for log in logs:
        result.append({
            "id": log.id,
            "user_msg": log.user_msg[:200],
            "risk_level": log.risk_level,
            "risk_alert": log.risk_alert[:300],
            "scene": log.scene,
            "create_time": log.create_time.strftime("%Y-%m-%d %H:%M:%S")
        })
    return jsonify(result)


@app.route('/stress-test', methods=['POST'])
def stress_test():
    global test_results
    test_results = {
        "total_requests": 0,
        "success_requests": 0,
        "fail_requests": 0,
        "avg_response_time": 0.0,
        "response_times": [],
        "error_details": defaultdict(int)
    }

    data = request.get_json()
    concurrency = int(data.get("concurrency", 10))
    total = int(data.get("total", 10))
    test_msg = data.get("test_msg", "劳动合同纠纷如何维权")

    def test_single_request():
        global test_results
        start_time = time.time()
        messages = [{"role": "system", "content": LEGAL_PROMPT_TEMPLATES["civil"]}]
        messages.append({"role": "user", "content": test_msg})
        result = call_deepseek(messages, "civil")
        end_time = time.time()
        response_time = end_time - start_time

        test_results["total_requests"] += 1
        if result["success"]:
            test_results["success_requests"] += 1
            test_results["response_times"].append(response_time)
        else:
            test_results["fail_requests"] += 1
            error = result.get("error", "未知错误")
            test_results["error_details"][error[:30]] += 1

        if test_results["response_times"]:
            test_results["avg_response_time"] = sum(test_results["response_times"]) / len(
                test_results["response_times"])

    threads = []
    for i in range(total):
        while len(threading.enumerate()) > concurrency:
            time.sleep(0.01)
        t = threading.Thread(target=test_single_request)
        threads.append(t)
        t.start()

    for t in threads:
        t.join()

    return jsonify(test_results)


@app.route('/test-result')
def get_test_result():
    return jsonify(test_results)


@app.route('/clear-test-result', methods=['POST'])
def clear_test_result():
    global test_results
    test_results = {
        "total_requests": 0,
        "success_requests": 0,
        "fail_requests": 0,
        "avg_response_time": 0.0,
        "response_times": [],
        "error_details": defaultdict(int)
    }
    return jsonify({"status": "cleared"})


# 初始化数据库
with app.app_context():
    db.create_all()

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5003)